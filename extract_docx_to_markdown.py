#!/usr/bin/env python3
"""Extract text from DOCX and convert to markdown."""

import sys
import os
from docx import Document

def convert_docx_to_markdown(docx_path, output_path):
    """Convert DOCX to markdown file."""
    print(f"Extracting text from: {docx_path}")
    
    if not os.path.exists(docx_path):
        print(f"Error: DOCX file not found at {docx_path}")
        return False
    
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Check if paragraph has heading style
                style_name = paragraph.style.name
                if 'Heading' in style_name:
                    level = style_name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        markdown_content.append('#' * level_num + ' ' + text)
                    except ValueError:
                        markdown_content.append(text)
                else:
                    markdown_content.append(text)
        
        # Handle tables
        for table in doc.tables:
            markdown_content.append('')  # Add blank line before table
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                markdown_content.append('| ' + ' | '.join(row_data) + ' |')
                # Add header separator after first row
                if i == 0:
                    markdown_content.append('| ' + ' | '.join(['---'] * len(row_data)) + ' |')
            markdown_content.append('')  # Add blank line after table
        
        # Write to markdown file
        with open(output_path, 'w', encoding='utf-8') as md_file:
            md_file.write('\n\n'.join(markdown_content))
        
        print(f"Successfully created: {output_path}")
        return True
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    docx_path = "2. Gauntlet Week 4/1. Artifacts/1. Requirements/New drops/[PLAT] Superbuilders - Educational Video Generation/PRD_Educational Video Generation.docx"
    md_path = "2. Gauntlet Week 4/1. Artifacts/1. Requirements/New drops/[PLAT] Superbuilders - Educational Video Generation/PRD_Educational Video Generation.md"
    
    convert_docx_to_markdown(docx_path, md_path)

